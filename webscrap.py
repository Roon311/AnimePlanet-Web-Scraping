import requests
import csv
from bs4 import BeautifulSoup
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import RGBColor

# will use lxml as the parser for beatifulsoup


def open_image_from_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  # Check for successful response

        image_data = BytesIO(response.content)
        img = Image.open(image_data)
        img.show()  # Open the image using the default image viewer
    except requests.exceptions.RequestException as e:
        print(f"Error requesting the image: {e}")
    except Exception as e:
        print(f"An error occurred: {e}")


season= input('Please input a season in the form season-YYYY: ')
if not season:
    season = 'summer-2023'
page =requests.get(f"https://www.anime-planet.com/anime/seasons/{season}")

def main(page):
    src=page.content
    soup=BeautifulSoup(src,"lxml")

    site_container=soup.find("div",id="siteContainer")
    doc = Document()

    # Add a title
    def get_anime_info(site_container):
        season_name=site_container.find("h1").text.strip()
        doc.add_heading(season_name, level=1)

        #The first container is the anime container, and the rest of the containers could be found through extracting h4 titles
        animetypes=['Series']
        temp=site_container.find_all("h4")
        temp = [tag.text.strip() for tag in temp]
        animetypes+=temp
        anime_containers=site_container.find_all("ul",{'class':'cardDeck'})
        #The number of the containers should match the number of anime types
        for typ,container in zip(animetypes,anime_containers):
            animes=container.find_all("li") # All animes for a specific type
            print(f'#########################{typ}#############################')
            print(f'###########################################################')
            doc.add_heading(typ, level=1)
            for anime in animes:
                number_of_episodes=anime.attrs['data-total-episodes']
                para=anime.find('a')
                description = para.contents[1].contents[0].attrs
                title=description['alt']
                imagelink=description['data-src']
                #open_image_from_url(image)

                description = para.contents[1].contents[0].attrs
                print(title)
                print(f'Number of episodes: {number_of_episodes}')
                doc.add_heading(title, level=2).bold = True
                doc.add_heading(number_of_episodes + " episodes", level=3).bold = False
                 # Add anime photo
                image_response = requests.get(imagelink)
                if image_response.status_code == 200:
                    image = Image.open(BytesIO(image_response.content))
                    image = image.convert("RGB")  # Convert to RGB format
                    image.thumbnail((156, 222))  # Resize the image
                    #width, height = image.size
                    image_path = "temp_image.jpg"
                    image.save(image_path)
                    
                    doc.add_picture(image_path)
                    image.close()
                print('...............................................')
                
        doc.save('anime_season.docx')

    get_anime_info(site_container)

main(page)
